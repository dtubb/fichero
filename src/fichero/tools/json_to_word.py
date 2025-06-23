#!/usr/bin/env python3
"""
JSON to Word converter
Converts JSON data from LLM processing to formatted Word documents
"""

import typer
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, Optional
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Support both standalone CLI usage and workflow executor imports
try:
    # When imported by workflow executor (absolute imports work)
    from fichero.tools.utils.batch import BatchProcessor
    from fichero.tools.utils.processor import process_file
    from fichero.tools.utils.segment_handler import SegmentHandler
    from fichero.tools.utils.files import ensure_dirs
    from fichero.tools.utils.manifest import ManifestProcessor
    from fichero.tools.utils.tool_logger import get_tool_logger
except ImportError:
    # When run as standalone script (relative imports work)
    from utils.batch import BatchProcessor
    from utils.processor import process_file
    from utils.segment_handler import SegmentHandler
    from utils.files import ensure_dirs
    from utils.manifest import ManifestProcessor
    from utils.tool_logger import get_tool_logger

tool_logger = get_tool_logger('json_to_word')

def set_document_properties(doc):
    """Set up basic document properties"""
    section = doc.sections[0]
    section.left_margin = section.right_margin = Inches(1)
    section.top_margin = section.bottom_margin = Inches(1)

def add_section_title(doc, title):
    """Add a Heading 2 section title in Helvetica"""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.name = 'Helvetica'
    run.font.size = Pt(16)
    run.font.bold = True
    p.space_after = Pt(8)
    return p

def add_title(doc, title):
    """Add document title as H1 heading"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.font.name = 'Helvetica'
    run.font.size = Pt(18)  # H1 size
    run.font.bold = True
    p.space_after = Pt(12)  # Reduced space after title

def add_section(doc, title, content):
    """Add a section with title and content"""
    # Add section title
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.name = 'Helvetica'
    run.font.size = Pt(14)
    run.font.bold = True
    p.space_after = Pt(6)
    
    # Add content
    p = doc.add_paragraph()
    run = p.add_run(content)
    run.font.name = 'Helvetica'
    run.font.size = Pt(11)
    p.space_after = Pt(12)

def set_cell_bottom_padding(cell, padding_pt=6):
    """Set bottom padding for a table cell in points (pt)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    bottom = tcMar.find(qn('w:bottom'))
    if bottom is None:
        bottom = OxmlElement('w:bottom')
        tcMar.append(bottom)
    bottom.set(qn('w:w'), str(int(padding_pt * 20)))  # 1pt = 20 twips
    bottom.set(qn('w:type'), 'dxa')

def add_table_section(doc, title, data_list):
    """Add a section with a table"""
    if not data_list or not isinstance(data_list, list):
        return
    
    # Add section title
    add_section_title(doc, title)
    
    # Get all unique keys for columns
    columns = set()
    for item in data_list:
        if isinstance(item, dict):
            columns.update(item.keys())
    columns = list(columns)
    
    if not columns:
        return
    
    # Create table
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = 'Table Grid'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col.replace('_', ' ').title()
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Helvetica'
                run.font.size = Pt(10)
                run.font.bold = True
            paragraph.space_after = Pt(6)
        set_cell_bottom_padding(hdr_cells[i], 6)
    
    # Add data rows
    for item in data_list:
        if isinstance(item, dict):
            row_cells = table.add_row().cells
            for i, col in enumerate(columns):
                val = item.get(col, "")
                if isinstance(val, list):
                    val = ", ".join(str(v) for v in val)
                elif isinstance(val, dict):
                    val = json.dumps(val, ensure_ascii=False)
                row_cells[i].text = str(val)
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Helvetica'
                        run.font.size = Pt(9)
                    paragraph.space_after = Pt(6)
                set_cell_bottom_padding(row_cells[i], 6)
    
    doc.add_paragraph()  # Add space after table

def convert_json_to_word(json_file_path: Path, output_path: Path, parent_folder: str) -> Dict:
    """Convert JSON file to Word document"""
    
    # Get relative path using SegmentHandler
    rel_path = SegmentHandler.get_relative_path(json_file_path)
    tool_logger.info(f"Processing JSON file: {rel_path}")
    
    try:
        # Read JSON file
        with open(json_file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        tool_logger.info(f"Successfully loaded JSON file with {len(data)} top-level keys")
        
        # Create new document
        doc = Document()
        set_document_properties(doc)
        tool_logger.info("Created new Word document with basic properties")
        
        # Add title (use passed parent folder name)
        add_title(doc, parent_folder)
        tool_logger.info(f"Added document title: {parent_folder}")
        
        # Get results section
        results = data.get("results", {})
        tool_logger.info(f"Found {len(results)} result sections to process")
        
        sections_added = 0
        
        # Define core sections that should come first
        core_sections = {
            "resumen": "Resumen",
            "summary": "Resumen",
            "personas_clave_y_etiquetas": "Personas Clave y Etiquetas",
            "key_people_and_tags": "Personas Clave y Etiquetas"
        }
        
        # Helper function to get field value with fallbacks
        def get_field_value(data, field_name):
            if isinstance(data, dict):
                if field_name == "resumen":
                    return data.get("resumen") or data.get("summary")
                elif field_name == "etiquetas":
                    return data.get("etiquetas") or data.get("tags")
                elif field_name == "personas_clave":
                    return data.get("personas_clave") or data.get("key_people")
                elif field_name == "nombre":
                    return data.get("nombre") or data.get("name")
                elif field_name == "contexto":
                    return data.get("contexto") or data.get("context")
            return None
        
        # Process core sections first
        for section_key, section_title in core_sections.items():
            if section_key in results:
                section_data = results[section_key]
                
                # Handle summary/resumen
                if section_key in ["resumen", "summary"]:
                    summary_text = get_field_value(section_data, "resumen")
                    if summary_text:
                        add_section(doc, section_title, summary_text)
                        sections_added += 1
                
                # Handle key people and tags
                elif section_key in ["personas_clave_y_etiquetas", "key_people_and_tags"]:
                    # Add tags/keywords
                    if isinstance(section_data, dict):
                        tags = get_field_value(section_data, "etiquetas")
                        if isinstance(tags, str):
                            tag_list = [tag.strip() for tag in tags.split(";") if tag.strip()]
                            tags_text = "; ".join(tag_list)
                            add_section(doc, "Palabras Clave", tags_text)
                            sections_added += 1
                        
                        # Add key people table
                        key_people = get_field_value(section_data, "personas_clave")
                        if isinstance(key_people, list):
                            key_people_rows = []
                            for item in key_people:
                                key_people_rows.append({
                                    "Nombre": get_field_value(item, "nombre") or "",
                                    "Contexto": get_field_value(item, "contexto") or ""
                                })
                            if key_people_rows:
                                add_table_section(doc, "Personas Clave", key_people_rows)
                                sections_added += 1
        
        # Process remaining sections in order
        for section_key, section_data in results.items():
            # Skip already processed core sections
            if section_key in core_sections:
                continue
            
            # Do NOT add a heading for the step/section here
            
            # Handle any section that contains a list of items
            if isinstance(section_data, dict):
                for category, items in section_data.items():
                    if isinstance(items, list) and items:
                        # Add a Heading 2 for the category
                        category_title = category.replace('_', ' ').title()
                        add_section_title(doc, category_title)
                        # Get all fields from the first item, preserving order
                        all_fields = list(items[0].keys())
                        rows = []
                        for item in items:
                            row = {}
                            for key in all_fields:
                                val = item.get(key, "")
                                if isinstance(val, list):
                                    val = ", ".join(str(v) for v in val)
                                elif isinstance(val, dict):
                                    val = json.dumps(val, ensure_ascii=False)
                                row[key] = str(val)
                            rows.append(row)
                        # Create and style table
                        table = doc.add_table(rows=1, cols=len(all_fields))
                        table.style = 'Table Grid'
                        # Add headers
                        hdr_cells = table.rows[0].cells
                        for i, col in enumerate(all_fields):
                            hdr_cells[i].text = col.replace('_', ' ').title()
                            for paragraph in hdr_cells[i].paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Helvetica'
                                    run.font.size = Pt(10)
                                    run.font.bold = True
                                paragraph.space_after = Pt(6)
                            set_cell_bottom_padding(hdr_cells[i], 6)
                        # Add data rows
                        for row in rows:
                            row_cells = table.add_row().cells
                            for i, col in enumerate(all_fields):
                                row_cells[i].text = str(row.get(col, ""))
                                for paragraph in row_cells[i].paragraphs:
                                    for run in paragraph.runs:
                                        run.font.name = 'Helvetica'
                                        run.font.size = Pt(9)
                                    paragraph.space_after = Pt(6)
                                set_cell_bottom_padding(row_cells[i], 6)
                        doc.add_paragraph()
                        sections_added += 1
        
        # Ensure output directory exists
        ensure_dirs(output_path)
        
        # Change extension to .docx
        output_path = output_path.with_suffix('.docx')
        
        # Save document
        doc.save(output_path)
        
        # Get relative path for output
        output_rel_path = SegmentHandler.get_relative_path(output_path)
        
        tool_logger.info(f"Successfully created Word document with {sections_added} sections")
        tool_logger.info(f"Output saved to: {output_rel_path}")
        
        return {
            "outputs": [str(output_rel_path)],
            "source": str(rel_path),
            "success": True,
            "details": {
                "sections_created": sections_added,
                "output_format": "docx"
            }
        }
        
    except FileNotFoundError:
        tool_logger.error(f"JSON file not found: {json_file_path}")
        return {
            "outputs": [],
            "source": str(rel_path),
            "success": False,
            "error": "File not found"
        }
    except json.JSONDecodeError as e:
        tool_logger.error(f"Invalid JSON in file {json_file_path}: {e}")
        return {
            "outputs": [],
            "source": str(rel_path),
            "success": False,
            "error": f"Invalid JSON: {e}"
        }
    except Exception as e:
        tool_logger.error(f"Error converting {json_file_path}: {e}")
        return {
            "outputs": [],
            "source": str(rel_path),
            "success": False,
            "error": str(e)
        }

def process_document(file_path: str, output_folder: Path) -> Dict:
    """Process a single JSON document file"""
    
    def process_fn(f: Path, o: Path) -> Dict:
        # Get relative path using SegmentHandler
        rel_path = SegmentHandler.get_relative_path(f)
        tool_logger.info(f"Processing document: {rel_path}")
        
        # Get the main folder name by going up 4 parents from the JSON file
        # Path structure: 1939-Fabriciano-Mosquera.../assets/llm_catalogue/documents/documents_summary.json
        # We need to go up: documents -> llm_catalogue -> assets -> main_folder
        file_path = Path(f)
        parent_folder = file_path.parent.parent.parent.parent.name
        tool_logger.info(f"Parent folder name: {parent_folder}")
        
        # Create output path using consistent path handling with catalogue suffix
        doc_output_path = output_folder / "documents" / rel_path.parent / f"{parent_folder}-catalogue.docx"
        ensure_dirs(doc_output_path)
        
        # Convert JSON to Word - pass parent_folder to ensure consistency
        return convert_json_to_word(f, doc_output_path, parent_folder)
    
    return process_file(
        file_path=file_path,
        output_folder=output_folder,
        process_fn=process_fn,
        file_types={'.json': process_fn}
    )

def json_to_word_batch(
    source_folder: Path,
    source_manifest: Path,
    output_folder: Path,
    **kwargs
) -> dict:
    """
    Convert catalogue JSON files to Word documents - importable function
    
    Args:
        source_folder: Source folder containing JSON files
        source_manifest: Manifest file from LLM processing
        output_folder: Output folder for Word documents
        
    Returns:
        Processing statistics dictionary
    """
    tool_logger.info("Converting JSON files to Word documents")
    tool_logger.info(f"Source folder: {source_folder}")
    tool_logger.info(f"Source manifest: {source_manifest}")
    tool_logger.info(f"Output folder: {output_folder}")
    
    # Ensure output directory exists
    ensure_dirs(output_folder)

    processor = BatchProcessor(
        input_manifest=source_manifest,
        output_folder=output_folder,
        process_name="json_to_word",
        base_folder=source_folder,
        processor_fn=lambda f, o: process_document(f, o),
        use_source=False
    )
    
    result = processor.process()
    
    tool_logger.success("Conversion completed!")
    tool_logger.success(f"Processed: {result.get('processed', 0)}")
    tool_logger.info(f"Skipped: {result.get('skipped', 0)}")
    tool_logger.error(f"Failed: {result.get('failed', 0)}")
    
    return result

def json_to_word(
    source_folder: Path = typer.Argument(..., help="Source folder containing JSON files"),
    source_manifest: Path = typer.Argument(..., help="Manifest file from LLM processing"),
    output_folder: Path = typer.Argument(..., help="Output folder for Word documents")
):
    """Convert JSON files from LLM processing to Word documents"""
    return json_to_word_batch(source_folder, source_manifest, output_folder)

if __name__ == "__main__":
    typer.run(json_to_word) 