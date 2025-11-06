#!/usr/bin/env python3
"""
Convert background-removed images and transcriptions to Word documents with side-by-side layout
"""

import typer
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_BREAK
from PIL import Image
import io
import json
import tempfile
import shutil
import subprocess
from datetime import datetime

# Support both standalone CLI usage and workflow executor imports
try:
    # When imported by workflow executor (absolute imports work)
    from fichero.tools.utils.batch import BatchProcessor
    from fichero.tools.utils.processor import process_file
    from fichero.tools.utils.segment_handler import SegmentHandler
    from fichero.tools.utils.files import ensure_dirs, get_relative_path
    from fichero.tools.utils.manifest import ManifestProcessor
    from fichero.tools.utils.tool_logger import get_tool_logger
except ImportError:
    # When run as standalone script (relative imports work)
    from fichero.tools.utils.batch import BatchProcessor
    from fichero.tools.utils.processor import process_file
    from fichero.tools.utils.segment_handler import SegmentHandler
    from fichero.tools.utils.files import ensure_dirs, get_relative_path
    from fichero.tools.utils.manifest import ManifestProcessor
    from fichero.tools.utils.tool_logger import get_tool_logger

tool_logger = get_tool_logger('convert_to_word')
app = typer.Typer()

# Update the grey color to a lighter shade
GREY_FILL = "E8E8E8"  # Lighter grey

class SpreadManager:
    """Manages temporary storage and processing of spreads"""
    def __init__(self, temp_dir: Path):
        self.temp_dir = temp_dir
        self.temp_dir.mkdir(parents=True, exist_ok=True)
        self.spreads = {}
        
    def add_spread(self, parent_folder: str, image_path: str, text: str, filename: str) -> bool:
        """Add a spread to the temporary storage"""
        if parent_folder not in self.spreads:
            self.spreads[parent_folder] = {
                "path": None,
                "processed": False,
                "spreads": [],
                "needs_rebuild": True
            }
            
            
        # Check if spread exists
        for existing_spread in self.spreads[parent_folder]["spreads"]:
            if existing_spread["image_path"] == image_path:
                if existing_spread["text"] != text:
                    existing_spread["text"] = text
                    self.spreads[parent_folder]["needs_rebuild"] = True
                return True
                
        # Add new spread
        self.spreads[parent_folder]["spreads"].append({
            "image_path": image_path,
            "text": text,
            "filename": filename
        })
        self.spreads[parent_folder]["needs_rebuild"] = True
        self.spreads[parent_folder]["processed"] = True
        return True
        
    def get_spreads(self, parent_folder: str) -> list:
        """Get spreads for a parent folder"""
        return self.spreads.get(parent_folder, {}).get("spreads", [])
        
    def needs_rebuild(self, parent_folder: str) -> bool:
        """Check if a document needs rebuilding"""
        return self.spreads.get(parent_folder, {}).get("needs_rebuild", False)
        
    def is_processed(self, parent_folder: str) -> bool:
        """Check if a document has been processed"""
        return self.spreads.get(parent_folder, {}).get("processed", False)
        
    def set_path(self, parent_folder: str, path: Path):
        """Set the path for a document"""
        if parent_folder in self.spreads:
            self.spreads[parent_folder]["path"] = path
            
    def get_path(self, parent_folder: str) -> Path:
        """Get the path for a document"""
        return self.spreads.get(parent_folder, {}).get("path")
        
    def cleanup(self):
        """Clean up temporary files"""
        if self.temp_dir.exists():
            shutil.rmtree(self.temp_dir)

def set_document_properties(doc):
    """Set up initial document properties"""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.different_first_page_header_footer = True
    # No default margins - we'll set them per section

def create_cover_page(doc, folder_name):
    """Create a white cover page with centered folder name as title"""
    # Create blank first page
    blank_section = doc.sections[0]
    blank_section.left_margin = blank_section.right_margin = blank_section.top_margin = blank_section.bottom_margin = Inches(1)
    doc.add_paragraph()  # Add empty paragraph for blank page
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    
    # Create new section for cover page
    cover_section = doc.add_section(WD_SECTION.NEW_PAGE)
    cover_section.left_margin = cover_section.right_margin = cover_section.top_margin = cover_section.bottom_margin = Inches(1)
    
    # Add vertical space to center content
    for _ in range(10):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Use the full folder name as title, preserving all special characters
    title = str(folder_name)  # Ensure string conversion
    run = p.add_run(title)
    run.font.name = 'Helvetica'
    run.font.size = Pt(24)
    run.font.bold = True
    
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def calculate_optimal_font_size(text_length, page_width_inches, page_height_inches):
    """Calculate optimal font size to fit text on one page"""
    # Approximate characters per line and lines per page at 12pt
    CHAR_WIDTH_12PT = 0.11
    LINE_HEIGHT_12PT = 0.18
    
    # Available space (accounting for margins)
    available_width = page_width_inches - 2
    available_height = page_height_inches - 2
    
    # Define font size ranges based on text length
    if text_length < 500:
        font_sizes = list(range(12, 10, -1))    # Short texts
    elif text_length < 1000:
        font_sizes = list(range(11, 9, -1))     # Medium texts
    else:
        font_sizes = list(range(10, 8, -1))     # Long texts
    
    for font_size in font_sizes:
        scale_factor = font_size / 12.0
        chars_per_line = int(available_width / (CHAR_WIDTH_12PT * scale_factor))
        lines_per_page = int(available_height / (LINE_HEIGHT_12PT * scale_factor))
        total_capacity = chars_per_line * lines_per_page
        
        if text_length <= total_capacity:
            return font_size
    
    return 8

def load_image_with_jxl_support(image_path: Path) -> Image.Image:
    """
    Load an image file, with support for JPEG XL format.
    Falls back to PIL for other formats.
    """
    if image_path.suffix.lower() == '.jxl':
        # Check if djxl is available
        if shutil.which('djxl'):
            try:
                # Convert JXL to temporary PNG using djxl
                temp_png = image_path.with_suffix('.temp.png')
                cmd = ['djxl', str(image_path), str(temp_png)]
                process = subprocess.run(cmd, capture_output=True, text=True)
                
                if process.returncode == 0 and temp_png.exists():
                    # Load the converted PNG
                    img = Image.open(temp_png)
                    # Clean up temp file
                    temp_png.unlink()
                    return img
                else:
                    tool_logger.warning(f"djxl conversion failed for {image_path}: {process.stderr}")
            except Exception as e:
                tool_logger.warning(f"Failed to convert JXL file {image_path}: {str(e)}")
        else:
            tool_logger.warning(f"djxl not installed, cannot read JXL file {image_path}")
            tool_logger.info("Install with: brew install libjxl (macOS) or apt-get install libjxl-tools (Ubuntu/Debian)")
        
        # If JXL conversion failed, raise an exception
        raise ValueError(f"Cannot read JXL file {image_path} - djxl not available or conversion failed")
    
    # For other formats, use PIL directly
    return Image.open(image_path)

def create_spread(doc, image_path, text, filename):
    """Create a two-page spread with image (left) and text (right)"""
    # Left page - Image
    section = doc.add_section()
    
    # Center image by setting equal margins
    margin = Inches(0.75)
    section.left_margin = section.right_margin = margin
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.75)
    
    # Add image paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    
    try:
        # Use the new function that supports JXL
        img = load_image_with_jxl_support(Path(image_path))
        
        # Convert to RGB (removing alpha channel)
        if img.mode in ('RGBA', 'LA') or (img.mode == 'P' and 'transparency' in img.info):
            bg = Image.new('RGB', img.size, 'white')
            if img.mode != 'RGBA':
                img = img.convert('RGBA')
            bg.paste(img, mask=img.split()[3])
            img = bg
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        
        # Calculate dimensions for page
        available_width = 8.5 - (2 * margin.inches)
        available_height = 9.75
        img_ratio = img.height / img.width
        
        if img_ratio > available_height/available_width:
            height = available_height
            width = height / img_ratio
        else:
            width = available_width
            height = width * img_ratio
        
        # Keep 300 DPI resolution for print quality
        target_width = int(width * 300)
        target_height = int(height * 300)
        img = img.resize((target_width, target_height), Image.Resampling.LANCZOS)
        
        # Save as JPG with compression
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format='JPEG', quality=80, optimize=True)
        img_byte_arr.seek(0)
        
        run.add_picture(img_byte_arr, width=Inches(width), height=Inches(height))
    except Exception as e:
        tool_logger.error(f"Error loading image {image_path}: {str(e)}")
        run.add_text(f"[Error loading image: {str(e)}]")

    # Right page - Text
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.left_margin = section.right_margin = Inches(1)
    section.top_margin = section.bottom_margin = Inches(1)
    
    # Add filename as title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run(filename)
    title_run.font.name = 'Helvetica'
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    title_p.space_after = Pt(12)  # Add space after title
    
    # Add text content
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Helvetica'
    run.font.size = Pt(10)
    
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    
    # Add page break for next spread
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    
    return True

def process_document(file_path: str, output_folder: Path, spread_manager: SpreadManager, transcription_folder: Path = None) -> dict:
    """Process a single document file using process_file utility"""
    file_path = Path(file_path)

    def process_fn(f: str, o: Path) -> dict:
        try:
            # Get relative path using SegmentHandler
            rel_path = SegmentHandler.get_relative_path(Path(f))
            tool_logger.info(f"Processing document: {f}")
            tool_logger.debug(f"Relative path: {rel_path}")

            # Get the parent folder name for document grouping
            parent_folder = Path(f).parent.name
            tool_logger.info(f"Parent folder name: {parent_folder}")

            # Find the corresponding text file using consistent path handling
            # If transcription_folder is provided, use it; otherwise default to transcriptions/documents
            if transcription_folder:
                text_path = transcription_folder / "documents" / rel_path.with_suffix('.txt')
            else:
                text_path = output_folder.parent / "transcriptions" / "documents" / rel_path.with_suffix('.txt')
            if not text_path.exists():
                # Try with raw path if not found
                text_path = Path(str(text_path).replace(';', '\\;'))
                if not text_path.exists():
                    tool_logger.error(f"Text file not found: {text_path}")
                    return {
                        "error": f"Text file not found: {text_path}",
                        "source": str(rel_path),
                        "success": False
                    }
            
            # Read the text content
            try:
                text = text_path.read_text(encoding='utf-8')
            except UnicodeDecodeError:
                # Fallback to system default encoding if UTF-8 fails
                text = text_path.read_text()
            
            # Check if the image file exists, try multiple extensions if not
            image_path = Path(f)
            if not image_path.exists():
                # Try different extensions in the same location
                for ext in ['.jpg', '.jpeg', '.png', '.jxl', '.tif', '.tiff']:
                    alt_path = image_path.with_suffix(ext)
                    if alt_path.exists():
                        tool_logger.info(f"Found image with different extension: {alt_path}")
                        image_path = alt_path
                        break
                else:
                    # If still not found, try looking in the base folder with different extensions
                    # This handles cases where the manifest has wrong extensions
                    base_path = Path(f).parent / Path(f).stem
                    for ext in ['.jpg', '.jpeg', '.png', '.jxl', '.tif', '.tiff']:
                        alt_path = base_path.with_suffix(ext)
                        if alt_path.exists():
                            tool_logger.info(f"Found image in base folder with different extension: {alt_path}")
                            image_path = alt_path
                            break
                    else:
                        # Last resort: try searching in the entire documents folder
                        documents_folder = Path(f).parent
                        stem = Path(f).stem
                        for ext in ['.jpg', '.jpeg', '.png', '.jxl', '.tif', '.tiff']:
                            search_pattern = f"{stem}{ext}"
                            matching_files = list(documents_folder.glob(search_pattern))
                            if matching_files:
                                tool_logger.info(f"Found image by searching: {matching_files[0]}")
                                image_path = matching_files[0]
                                break
                        else:
                            tool_logger.error(f"Image file not found: {f}")
                            return {
                                "error": f"Image file not found: {f}",
                                "source": str(rel_path),
                                "success": False
                            }
            
            # Add spread to manager
            spread_manager.add_spread(parent_folder, str(image_path), text, Path(f).stem)
            
            # Get the final document path using consistent path handling
            doc_output_path = output_folder / "documents" / rel_path.parent / f"{parent_folder}.docx"
            spread_manager.set_path(parent_folder, doc_output_path)
            
            return {
                "outputs": [str(SegmentHandler.get_relative_path(doc_output_path))],
                "source": str(rel_path),
                "details": {
                    "text_length": len(text),
                    "needs_rebuild": spread_manager.needs_rebuild(parent_folder),
                    "spread_count": len(spread_manager.get_spreads(parent_folder))
                }
            }
            
        except Exception as e:
            tool_logger.error(f"Error processing {f}: {str(e)}")
            return {
                "error": str(e),
                "source": str(rel_path)
            }
    
    return process_file(
        file_path=str(file_path),
        output_folder=output_folder,
        process_fn=lambda f, o: process_fn(f, o),
        file_types={
            '.png': process_fn,
            '.jpg': process_fn,
            '.jpeg': process_fn,
            '.jxl': process_fn,  # Add JXL support
            '.tif': process_fn,  # Add TIFF support
            '.tiff': process_fn  # Add TIFF support
        }
    )

def convert_to_word_batch(
    images_folder: Path,
    transcription_manifest: Path,
    output_folder: Path,
    transcription_folder: Path = None,
    skip_processing: bool = False,
    **kwargs
) -> dict:
    """
    Convert background-removed images and transcriptions to Word documents with side-by-side layout - importable function

    Args:
        images_folder: Input background removed images folder
        transcription_manifest: Input transcription manifest
        output_folder: Output folder for Word documents
        transcription_folder: Optional folder containing transcription text files (defaults to manifest's parent folder)
        skip_processing: If True, create empty Word files for fast testing (default: False)

    Returns:
        Processing statistics dictionary
    """
    tool_logger.info(f"Converting images in {images_folder} to Word documents")
    if skip_processing:
        tool_logger.info("⚡ SKIP MODE: Creating empty Word files for testing")

        # Read manifest to get list of items to process
        manifest_entries = []
        if Path(transcription_manifest).exists():
            with open(transcription_manifest, 'r') as f:
                for line in f:
                    if line.strip():
                        manifest_entries.append(json.loads(line))

        # Create output folder
        output_folder = Path(output_folder)
        output_folder.mkdir(parents=True, exist_ok=True)

        # Create output manifest
        output_manifest_path = output_folder / "convert_to_word_manifest.jsonl"

        # Group entries by parent folder to create one Word doc per folder
        folder_groups = {}
        for entry in manifest_entries:
            # Get the source or output path
            source = entry.get('source', entry.get('outputs', [None])[0] if entry.get('outputs') else None)
            if not source:
                continue

            # Extract parent folder name (e.g., "1931 Antonio Asprilla.../file.txt" -> "1931 Antonio Asprilla...")
            parent = str(Path(source).parent)
            if parent == '.':
                parent = Path(source).stem  # Use filename without extension if no folder

            if parent not in folder_groups:
                folder_groups[parent] = []
            folder_groups[parent].append(entry)

        # Create empty Word documents for each folder
        stats = {
            'total': len(manifest_entries),
            'success': 0,
            'failed': 0,
            'skipped': 0
        }

        with open(output_manifest_path, 'w') as manifest_file:
            for parent_folder, entries in folder_groups.items():
                try:
                    # Create output path
                    output_path = output_folder / parent_folder / f"{Path(parent_folder).name}.docx"
                    output_path.parent.mkdir(parents=True, exist_ok=True)

                    # Create empty document
                    doc = Document()
                    doc.save(str(output_path))

                    # Write manifest entries for all items in this folder
                    for entry in entries:
                        source = entry.get('source', entry.get('outputs', [None])[0] if entry.get('outputs') else None)
                        manifest_entry = {
                            "source": source,
                            "outputs": [str(output_path.relative_to(output_folder))],
                            "processed_at": datetime.now().isoformat(),
                            "success": True,
                            "details": {"skip_processing": True}
                        }
                        manifest_file.write(json.dumps(manifest_entry) + "\n")
                        stats['success'] += 1

                    tool_logger.info(f"Created empty Word file: {output_path}")

                except Exception as e:
                    tool_logger.error(f"Error creating empty Word file for {parent_folder}: {str(e)}")
                    for entry in entries:
                        source = entry.get('source', entry.get('outputs', [None])[0] if entry.get('outputs') else None)
                        manifest_entry = {
                            "source": source,
                            "outputs": [],
                            "processed_at": datetime.now().isoformat(),
                            "success": False,
                            "error": str(e)
                        }
                        manifest_file.write(json.dumps(manifest_entry) + "\n")
                        stats['failed'] += 1

        tool_logger.success(f"⚡ Skip mode complete: {stats['success']} empty Word files created")
        return stats

    # Use provided transcription folder, or derive from manifest path
    # E.g., if manifest is "assets/enhanced/enhance_manifest.jsonl", folder is "assets/enhanced"
    # But for segmented workflows, we might want transcriptions from "assets/cleaned"
    if transcription_folder is None:
        transcription_folder = Path(transcription_manifest).parent
    else:
        transcription_folder = Path(transcription_folder)
    tool_logger.info(f"Using transcription folder: {transcription_folder}")

    # Create temporary directory for spreads
    with tempfile.TemporaryDirectory() as temp_dir:
        spread_manager = SpreadManager(Path(temp_dir))

        # Create processor for image files
        # Use outputs from the images manifest (not source, which points to originals)
        # We need to match images from images_folder with transcriptions from transcription_folder
        processor = BatchProcessor(
            input_manifest=transcription_manifest,
            output_folder=output_folder,
            process_name="convert_to_word",
            base_folder=images_folder,
            processor_fn=lambda f, o: process_document(f, o, spread_manager, transcription_folder),
            use_source=False,  # Use "outputs" field which contains actual processed images
            add_documents_prefix=False  # Images folder structure preserved from source
        )
        
        results = processor.process()
        
        # Process accumulated spreads and save complete documents
        for parent_folder in spread_manager.spreads:
            if not spread_manager.is_processed(parent_folder):
                tool_logger.warning(f"Document for folder {parent_folder} was not processed")
                continue
                
            try:
                # Only rebuild if needed
                if spread_manager.needs_rebuild(parent_folder):
                    # Create new document
                    doc = Document()
                    set_document_properties(doc)
                    create_cover_page(doc, parent_folder)
                    
                    # Add all spreads to the document
                    for spread in spread_manager.get_spreads(parent_folder):
                        create_spread(doc, spread["image_path"], spread["text"], spread["filename"])
                    
                    # Save complete document using consistent path handling
                    out_path = spread_manager.get_path(parent_folder)
                    ensure_dirs(out_path)
                    doc.save(str(out_path))
                    tool_logger.success(f"Saved complete document: {out_path}")
                else:
                    tool_logger.info(f"Skipping unchanged document: {parent_folder}")
            except Exception as e:
                tool_logger.error(f"Error saving document for {parent_folder}: {str(e)}")
        
        return results

@app.command()
def convert_to_word(
    background_removed_folder: Path = typer.Argument(..., help="Input background removed images folder"),
    transcription_manifest: Path = typer.Argument(..., help="Input transcription manifest"),
    word_folder: Path = typer.Argument(..., help="Output folder for Word documents")
):
    """Convert background-removed images and transcriptions to Word documents with side-by-side layout"""
    return convert_to_word_batch(background_removed_folder, transcription_manifest, word_folder)

if __name__ == "__main__":
    app()