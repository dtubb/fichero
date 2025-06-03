#!/usr/bin/env python3
"""
Simple JSON to Word converter
Converts JSON data from LLM processing to a formatted Word document
"""

import json
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_document_properties(doc):
    """Set up basic document properties"""
    section = doc.sections[0]
    section.left_margin = section.right_margin = Inches(1)
    section.top_margin = section.bottom_margin = Inches(1)

def add_title(doc, title):
    """Add document title"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = 'Arial'
    run.font.size = Pt(24)
    run.font.bold = True
    p.space_after = Pt(18)

def add_section(doc, title, content):
    """Add a section with title and content"""
    # Add section title
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    p.space_after = Pt(6)
    
    # Add content
    p = doc.add_paragraph()
    run = p.add_run(content)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    p.space_after = Pt(12)

def add_table_section(doc, title, data_list):
    """Add a section with a table"""
    if not data_list or not isinstance(data_list, list):
        return
    
    # Add section title
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    p.space_after = Pt(6)
    
    # Get all unique keys for columns
    columns = set()
    for item in data_list:
        if isinstance(item, dict):
            columns.update(item.keys())
    columns = list(columns)
    
    if not columns:
        return
    
    # Create table
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = 'Table Grid'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col.replace('_', ' ').title()
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.font.bold = True
    
    # Add data rows
    for item in data_list:
        if isinstance(item, dict):
            row_cells = table.add_row().cells
            for i, col in enumerate(columns):
                val = item.get(col, "")
                if isinstance(val, list):
                    val = ", ".join(str(v) for v in val)
                elif isinstance(val, dict):
                    val = json.dumps(val, ensure_ascii=False)
                row_cells[i].text = str(val)
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(9)
    
    doc.add_paragraph()  # Add space after table

def convert_json_to_word(json_file_path, output_path):
    """Convert JSON file to Word document"""
    
    # Read JSON file
    with open(json_file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Create new document
    doc = Document()
    set_document_properties(doc)
    
    # Add title (use folder name or default)
    folder_name = Path(json_file_path).parent.name
    add_title(doc, folder_name)
    
    # Get results section
    results = data.get("results", {})
    
    # Helper function to parse JSON strings
    def try_parse_json(val):
        if isinstance(val, str):
            try:
                return json.loads(val)
            except json.JSONDecodeError:
                return val
        return val
    
    # 1. Add Summary
    if "summary" in results:
        summary_data = results["summary"]
        if isinstance(summary_data, dict) and "summary" in summary_data:
            add_section(doc, "Summary", summary_data["summary"])
        elif isinstance(summary_data, str):
            add_section(doc, "Summary", summary_data)
    
    # 2. Add Keywords as a heading and paragraph
    if "key_people_and_tags" in results:
        key_people_tags_data = results["key_people_and_tags"]
        if isinstance(key_people_tags_data, dict) and "tags" in key_people_tags_data:
            tags = key_people_tags_data["tags"]
            if isinstance(tags, str):
                tag_list = [tag.strip() for tag in tags.split(";") if tag.strip()]
                tags_text = "; ".join(tag_list)
                # Add as a heading and paragraph
                p = doc.add_paragraph()
                run = p.add_run("Keywords")
                run.font.name = 'Arial'
                run.font.size = Pt(14)
                run.font.bold = True
                p.space_after = Pt(6)
                p = doc.add_paragraph(tags_text)
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                p.space_after = Pt(12)
    # 3. Add Key People (Name, Context) - always immediately after Keywords
    if "key_people_and_tags" in results:
        key_people_tags_data = results["key_people_and_tags"]
        if isinstance(key_people_tags_data, dict):
            if "key_people" in key_people_tags_data and isinstance(key_people_tags_data["key_people"], list):
                key_people_rows = []
                for item in key_people_tags_data["key_people"]:
                    key_people_rows.append({
                        "Name": item.get("name", ""),
                        "Context": item.get("context", "")
                    })
                if key_people_rows:
                    # Add 'Key People' heading
                    p = doc.add_paragraph()
                    run = p.add_run("Key People")
                    run.font.name = 'Arial'
                    run.font.size = Pt(14)
                    run.font.bold = True
                    p.space_after = Pt(6)
                    # Add table with explicit column order
                    columns = ["Name", "Context"]
                    table = doc.add_table(rows=1, cols=len(columns))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for i, col in enumerate(columns):
                        hdr_cells[i].text = col
                        for paragraph in hdr_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Arial'
                                run.font.size = Pt(10)
                                run.font.bold = True
                    for row in key_people_rows:
                        row_cells = table.add_row().cells
                        for i, col in enumerate(columns):
                            row_cells[i].text = str(row.get(col, ""))
                            for paragraph in row_cells[i].paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(9)
                    doc.add_paragraph()
    
    # 4. Add Timeline
    if "timeline_events" in results:
        timeline_data = results["timeline_events"]
        if isinstance(timeline_data, dict) and "timeline" in timeline_data:
            timeline = timeline_data["timeline"]
            if isinstance(timeline, list) and timeline:
                add_table_section(doc, "Timeline", timeline)
    
    # 5. Process NER data from multiple steps
    ner_data = {}
    ner_step_names = [
        "extract_ner_people_orgs_locations",
        "extract_ner_dates_legal_rivers", 
        "extract_ner_specialized_entities"
    ]
    for step_name in ner_step_names:
        if step_name in results:
            step_data = results[step_name]
            if isinstance(step_data, dict):
                ner_data.update(step_data)
            elif isinstance(step_data, str):
                try:
                    parsed_step_data = json.loads(step_data)
                    if isinstance(parsed_step_data, dict):
                        ner_data.update(parsed_step_data)
                except json.JSONDecodeError:
                    continue
    # Add NER tables with explicit column order, no sorting
    if ner_data:
        # Persons, Organizations, Locations: Name, Alternative Spellings, Context
        for category in ["persons", "organizations", "locations"]:
            if category in ner_data and isinstance(ner_data[category], list) and ner_data[category]:
                rows = []
                for item in ner_data[category]:
                    rows.append({
                        "Name": item.get("name", ""),
                        "Alternative Spellings": ", ".join(item.get("alternative_spellings", [])) if isinstance(item.get("alternative_spellings", []), list) else item.get("alternative_spellings", ""),
                        "Context": item.get("context", "")
                    })
                columns = ["Name", "Alternative Spellings", "Context"]
                table = doc.add_table(rows=1, cols=len(columns))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(columns):
                    hdr_cells[i].text = col
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)
                            run.font.bold = True
                for row in rows:
                    row_cells = table.add_row().cells
                    for i, col in enumerate(columns):
                        row_cells[i].text = str(row.get(col, ""))
                        for paragraph in row_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Arial'
                                run.font.size = Pt(9)
                doc.add_paragraph()
        # Dates, Legal Refs, Rivers: Name, Context
        for category in ["dates", "legal_refs", "rivers"]:
            if category in ner_data and isinstance(ner_data[category], list) and ner_data[category]:
                rows = []
                for item in ner_data[category]:
                    rows.append({
                        "Name": item.get("name", item.get("date", "")),
                        "Context": item.get("context", "")
                    })
                columns = ["Name", "Context"]
                table = doc.add_table(rows=1, cols=len(columns))
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(columns):
                    hdr_cells[i].text = col
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)
                            run.font.bold = True
                for row in rows:
                    row_cells = table.add_row().cells
                    for i, col in enumerate(columns):
                        row_cells[i].text = str(row.get(col, ""))
                        for paragraph in row_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Arial'
                                run.font.size = Pt(9)
                doc.add_paragraph()
        # All other categories: keep default behavior
        for category in ner_data:
            if category not in ["persons", "organizations", "locations", "dates", "legal_refs", "rivers"]:
                if isinstance(ner_data[category], list) and ner_data[category]:
                    add_table_section(doc, category.replace('_', ' ').title(), ner_data[category])
    
    # Save document
    doc.save(output_path)
    print(f"Word document saved to: {output_path}")

def main():
    if len(sys.argv) != 3:
        print("Usage: python simple_json_to_word.py <input_json_file> <output_word_file>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    if not Path(input_file).exists():
        print(f"Error: Input file '{input_file}' does not exist")
        sys.exit(1)
    
    try:
        convert_json_to_word(input_file, output_file)
        print("Conversion completed successfully!")
    except Exception as e:
        print(f"Error during conversion: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main() 