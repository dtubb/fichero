import typer
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_BREAK
from rich.console import Console
import json
import srsly
import re
import ast
import logging
from datetime import datetime
from typing import Any, Dict, List, Union

from utils.batch import BatchProcessor
from utils.segment_handler import SegmentHandler
from utils.files import ensure_dirs, get_relative_path
from utils.processor import process_file

# Initialize Typer app
app = typer.Typer()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
console = Console()

def set_document_properties(doc):
    """Set up initial document properties"""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = section.right_margin = Inches(1)
    section.top_margin = section.bottom_margin = Inches(1)

def add_title(doc, title):
    """Add the document title at the top of the document."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = 'Helvetica Neue'
    run.font.size = Pt(24)
    run.font.bold = True
    p.space_after = Pt(18)

def add_section(doc, title, content):
    """Add a section with title and content, no page break."""
    # Add section title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.font.name = 'Helvetica Neue'
    run.font.size = Pt(12)
    run.font.bold = True
    p.space_after = Pt(6)
    
    # Add content
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(content)
    run.font.name = 'Helvetica Neue'
    run.font.size = Pt(10)
    
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.15

def add_table_section(doc, title, data_list):
    """Add a section with a title and a table for a list of dicts."""
    if not data_list:
        return
    # Add section title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.font.name = 'Helvetica Neue'
    run.font.size = Pt(12)
    run.font.bold = True
    p.space_after = Pt(6)
    # Get all unique keys for columns
    columns = set()
    for item in data_list:
        if isinstance(item, dict):
            columns.update(item.keys())
    columns = list(columns)
    if not columns:
        return
    # Add table
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col.replace('_', ' ').title()
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Helvetica Neue'
                run.font.size = Pt(10)
    for item in data_list:
        row_cells = table.add_row().cells
        for i, col in enumerate(columns):
            val = item.get(col, "")
            if isinstance(val, list):
                val = ", ".join(str(v) for v in val)
            elif isinstance(val, dict):
                val = json.dumps(val, ensure_ascii=False)
            row_cells[i].text = str(val)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Helvetica Neue'
                    run.font.size = Pt(10)
    # Set narrow column width for Dates table
    if title.strip().lower() == 'dates' and len(columns) > 0:
        table.columns[0].width = Inches(1.2)
    doc.add_paragraph()  # Add space after table

def format_value(value: Any, indent: int = 0) -> str:
    """Format any value into a readable string with proper indentation"""
    if value is None:
        return "None"
    
    # Handle strings that might be JSON
    if isinstance(value, str):
        try:
            # Try to parse as JSON
            parsed = json.loads(value)
            return format_value(parsed, indent)
        except json.JSONDecodeError:
            # If not JSON, check for markdown code blocks
            if "```" in value:
                # Extract content from code blocks
                blocks = re.findall(r"```(?:json)?(.*?)```", value, re.DOTALL)
                if blocks:
                    try:
                        parsed = json.loads(blocks[0].strip())
                        return format_value(parsed, indent)
                    except json.JSONDecodeError:
                        pass
            return value
    
    # Handle dictionaries
    if isinstance(value, dict):
        if not value:
            return "{}"
        lines = []
        for k, v in value.items():
            # Format key nicely
            if isinstance(k, str):
                # Convert snake_case or camelCase to Title Case
                k = re.sub(r'[_\-]', ' ', k)
                k = re.sub(r'([a-z])([A-Z])', r'\1 \2', k)
                k = k.title()
            # Format value with proper indentation
            v_str = format_value(v, indent + 2)
            if isinstance(v, (dict, list)):
                lines.append(f"{' ' * indent}• {k}:")
                lines.append(v_str)
            else:
                lines.append(f"{' ' * indent}• {k}: {v_str}")
        return "\n".join(lines)
    
    # Handle lists
    if isinstance(value, list):
        if not value:
            return "[]"
        lines = []
        for item in value:
            item_str = format_value(item, indent + 2)
            if isinstance(item, (dict, list)):
                lines.append(f"{' ' * indent}•")
                lines.append(item_str)
            else:
                lines.append(f"{' ' * indent}• {item_str}")
        return "\n".join(lines)
    
    # Handle other types
    return str(value)

def detect_content_type(data: Dict) -> str:
    """Detect the type of content based on structure and keys"""
    # Check for common patterns
    if "persons" in data or "organizations" in data or "locations" in data:
        return "ner"
    if any(k.startswith("dc:") for k in data.keys()):
        return "dublin_core"
    if "summary" in data or "abstract" in data:
        return "summary"
    if "steps" in data and "results" in data:
        return "pipeline"
    return "generic"

def format_content(data: Dict) -> str:
    """Format content based on its detected type"""
    content_type = detect_content_type(data)
    
    if content_type == "ner":
        return format_ner_content(data)
    elif content_type == "dublin_core":
        return format_dublin_core_content(data)
    elif content_type == "summary":
        return format_summary_content(data)
    elif content_type == "pipeline":
        return format_pipeline_content(data)
    else:
        return format_value(data)

def format_ner_content(data: Dict) -> str:
    """Format NER content with entities and their occurrences"""
    sections = []
    for entity_type in ["persons", "organizations", "locations", "dates"]:
        if entity_type in data:
            entities = data[entity_type]
            if entities:
                sections.append(f"{entity_type.title()}:")
                for entity, pages in entities.items():
                    sections.append(f"  • {entity} (pages {', '.join(map(str, pages))})")
                sections.append("")
    return "\n".join(sections)

def format_dublin_core_content(data: Dict) -> str:
    """Format Dublin Core metadata"""
    sections = []
    field_order = [
        ('dc:title', 'Title'),
        ('dc:creator', 'Creator'),
        ('dc:contributor', 'Contributors'),
        ('dc:date', 'Date'),
        ('dc:type', 'Type'),
        ('dc:format', 'Format'),
        ('dc:identifier', 'Identifier'),
        ('dc:source', 'Source'),
        ('dc:language', 'Language'),
        ('dc:relation', 'Related Entity'),
        ('dc:coverage', 'Coverage'),
        ('dc:rights', 'Rights'),
        ('dc:description', 'Description'),
        ('dc:subject', 'Subjects')
    ]
    
    for dc_key, label in field_order:
        if dc_key in data:
            value = data[dc_key]
            if isinstance(value, list):
                sections.append(f"{label}:")
                for item in value:
                    if isinstance(item, dict):
                        if 'name' in item and 'role' in item:
                            sections.append(f"  • {item['name']} ({item['role']})")
                        else:
                            sections.append(f"  • {json.dumps(item)}")
                    else:
                        sections.append(f"  • {item}")
            else:
                sections.append(f"{label}: {value}")
            sections.append("")
    return "\n".join(sections)

def format_summary_content(data: Dict) -> str:
    """Format summary content with key points"""
    sections = []
    if "summary" in data:
        sections.append("Summary:")
        sections.append(data["summary"])
        sections.append("")
    if "key_points" in data:
        sections.append("Key Points:")
        for point in data["key_points"]:
            sections.append(f"• {point}")
    return "\n".join(sections)

def format_pipeline_content(data: Dict) -> str:
    """Format pipeline results with steps and their outputs"""
    sections = []
    if "steps" in data:
        sections.append("Processing Steps:")
        for step in data["steps"]:
            sections.append(f"• {step}")
        sections.append("")
    if "results" in data:
        sections.append("Results:")
        for step_name, result in data["results"].items():
            sections.append(f"\n{step_name.replace('_', ' ').title()}:")
            sections.append(format_value(result, 2))
    return "\n".join(sections)

def process_document(file_path: str, output_folder: Path) -> dict:
    """Process a single document file"""
    file_path = Path(file_path)
    
    def process_fn(f: Path, o: Path) -> dict:
        try:
            # Read summary file
            with open(f, 'r', encoding='utf-8') as fin:
                data = json.load(fin)
            
            # Get folder name from the file path
            folder_name = f.parent.name
            
            # Create new document
            doc = Document()
            set_document_properties(doc)
            add_title(doc, folder_name)
            
            # Add metadata section if available
            if "metadata" in data:
                add_section(doc, "Metadata", format_value(data["metadata"]))
            
            # Process results
            results = data.get("results", data)
            
            # Helper function to parse JSON strings
            def try_parse_json(val):
                if isinstance(val, str):
                    try:
                        return json.loads(val)
                    except json.JSONDecodeError:
                        return val
                return val
            
            # 1. Summary
            if "summary" in results:
                summary_data = results["summary"]
                if isinstance(summary_data, dict) and "summary" in summary_data:
                    add_section(doc, "Summary", summary_data["summary"])
                elif isinstance(summary_data, str):
                    add_section(doc, "Summary", summary_data)
            
            # 2. Timeline
            if "timeline_events" in results:
                timeline_data = results["timeline_events"]
                if isinstance(timeline_data, dict) and "timeline" in timeline_data:
                    timeline = timeline_data["timeline"]
                    if isinstance(timeline, list) and timeline and isinstance(timeline[0], dict):
                        add_table_section(doc, "Timeline", timeline)
            
            # 3. NER data processing
            table_keys = ["persons", "organizations", "locations", "dates", "key_events", "legal_refs"]
            
            if "extract_ner_multipage" in results:
                ner_data = try_parse_json(results["extract_ner_multipage"])
                if isinstance(ner_data, dict):
                    for key in table_keys:
                        if key in ner_data and isinstance(ner_data[key], list) and ner_data[key]:
                            if isinstance(ner_data[key][0], dict):
                                add_table_section(doc, key.replace('_', ' ').title(), ner_data[key])
                else:
                    add_section(doc, "NER Export Error", "NER data could not be parsed.")
            
            # 4. Process other table data directly from results
            for key, value in results.items():
                if key not in ["summary", "timeline_events", "extract_ner_multipage"]:
                    parsed_value = try_parse_json(value)
                    if key in table_keys and isinstance(parsed_value, list) and parsed_value:
                        if isinstance(parsed_value[0], dict):
                            add_table_section(doc, key.replace('_', ' ').title(), parsed_value)

            # Save document
            output_path = o / "documents-summary.docx"
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            logger.info(f"Saved document: {output_path}")
            
            # Return manifest-compatible output
            return {
                "outputs": [str(output_path.relative_to(output_folder))],
                "source": str(f.relative_to(f.parent.parent)),
                "type": "document",
                "details": {
                    "has_content": True,
                    "document_type": "word",
                    "content_type": detect_content_type(data),
                    "processed_at": datetime.now().isoformat()
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing {f}: {str(e)}")
            import traceback
            logger.error(f"Traceback: {traceback.format_exc()}")
            return {
                "error": str(e),
                "source": str(f.relative_to(f.parent.parent)),
                "type": "document"
            }
    
    return process_file(
        file_path=str(file_path),
        output_folder=output_folder,
        process_fn=process_fn
    )

@app.command()
def process_file_cmd(
    input_file: Path = typer.Argument(..., help="Input summary JSON file"),
    output_folder: Path = typer.Argument(..., help="Output folder for Word document")
):
    """Convert a single LLM output JSON file to a formatted Word document"""
    process_document(input_file, output_folder)

@app.command()
def process_folder(
    input_folder: Path = typer.Argument(..., help="Input folder containing LLM output JSONL files"),
    output_folder: Path = typer.Argument(..., help="Output folder for Word documents")
):
    """Convert LLM output JSONL files to formatted Word documents"""
    console.print(f"[green]Converting LLM output in {input_folder} to Word documents")
    
    # Create output folder if it doesn't exist
    output_folder.mkdir(parents=True, exist_ok=True)
    
    # Use BatchProcessor for consistent file handling
    processor = BatchProcessor(
        input_manifest=input_folder / "llm_process_manifest.jsonl",
        output_folder=output_folder,
        process_name="llm_to_word",
        processor_fn=process_document,
        base_folder=input_folder,
        use_source=True  # Use source paths from manifest
    )
    
    return processor.process()

if __name__ == "__main__":
    app() 