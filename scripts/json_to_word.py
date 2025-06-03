#!/usr/bin/env python3
"""
JSON to Word converter
Converts JSON data from LLM processing to formatted Word documents
"""

import typer
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from rich.console import Console
import logging
from typing import Dict, Any, Optional

# Import utility modules
from utils.batch import BatchProcessor
from utils.processor import process_file
from utils.segment_handler import SegmentHandler
from utils.files import ensure_dirs
from utils.manifest import ManifestProcessor

# Configure logging
logging.basicConfig(level=logging.WARNING)
logger = logging.getLogger(__name__)
console = Console()

def set_document_properties(doc):
    """Set up basic document properties"""
    section = doc.sections[0]
    section.left_margin = section.right_margin = Inches(1)
    section.top_margin = section.bottom_margin = Inches(1)

def add_title(doc, title):
    """Add document title"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = 'Arial'
    run.font.size = Pt(24)
    run.font.bold = True
    p.space_after = Pt(18)

def add_section(doc, title, content):
    """Add a section with title and content"""
    # Add section title
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    p.space_after = Pt(6)
    
    # Add content
    p = doc.add_paragraph()
    run = p.add_run(content)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    p.space_after = Pt(12)

def add_table_section(doc, title, data_list):
    """Add a section with a table"""
    if not data_list or not isinstance(data_list, list):
        return
    
    # Add section title
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.name = 'Arial'
    run.font.size = Pt(14)
    run.font.bold = True
    p.space_after = Pt(6)
    
    # Get all unique keys for columns
    columns = set()
    for item in data_list:
        if isinstance(item, dict):
            columns.update(item.keys())
    columns = list(columns)
    
    if not columns:
        return
    
    # Create table
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = 'Table Grid'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col.replace('_', ' ').title()
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.font.bold = True
    
    # Add data rows
    for item in data_list:
        if isinstance(item, dict):
            row_cells = table.add_row().cells
            for i, col in enumerate(columns):
                val = item.get(col, "")
                if isinstance(val, list):
                    val = ", ".join(str(v) for v in val)
                elif isinstance(val, dict):
                    val = json.dumps(val, ensure_ascii=False)
                row_cells[i].text = str(val)
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(9)
    
    doc.add_paragraph()  # Add space after table

def convert_json_to_word(json_file_path: Path, output_path: Path) -> Dict:
    """Convert JSON file to Word document"""
    
    # Get relative path using SegmentHandler
    rel_path = SegmentHandler.get_relative_path(json_file_path)
    logger.info(f"Processing JSON file: {rel_path}")
    
    try:
        # Read JSON file
        with open(json_file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        logger.info(f"Successfully loaded JSON file with {len(data)} top-level keys")
        
        # Create new document
        doc = Document()
        set_document_properties(doc)
        logger.info("Created new Word document with basic properties")
        
        # Add title (use folder name or default)
        folder_name = json_file_path.parent.name
        add_title(doc, folder_name)
        logger.info(f"Added document title: {folder_name}")
        
        # Get results section
        results = data.get("results", {})
        logger.info(f"Found {len(results)} result sections to process")
        
        # Helper function to parse JSON strings
        def try_parse_json(val):
            if isinstance(val, str):
                try:
                    return json.loads(val)
                except json.JSONDecodeError:
                    return val
            return val
        
        sections_added = 0
        
        # 1. Add Summary
        if "summary" in results:
            summary_data = results["summary"]
            if isinstance(summary_data, dict) and "summary" in summary_data:
                add_section(doc, "Summary", summary_data["summary"])
                sections_added += 1
                logger.info("Added Summary section from dict")
            elif isinstance(summary_data, str):
                add_section(doc, "Summary", summary_data)
                sections_added += 1
                logger.info("Added Summary section from string")
        
        # 2. Add Keywords as a heading and paragraph
        if "key_people_and_tags" in results:
            key_people_tags_data = results["key_people_and_tags"]
            if isinstance(key_people_tags_data, dict) and "tags" in key_people_tags_data:
                tags = key_people_tags_data["tags"]
                if isinstance(tags, str):
                    tag_list = [tag.strip() for tag in tags.split(";") if tag.strip()]
                    tags_text = "; ".join(tag_list)
                    # Add as a heading and paragraph
                    p = doc.add_paragraph()
                    run = p.add_run("Keywords")
                    run.font.name = 'Arial'
                    run.font.size = Pt(14)
                    run.font.bold = True
                    p.space_after = Pt(6)
                    p = doc.add_paragraph(tags_text)
                    for run in p.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(11)
                    p.space_after = Pt(12)
                    sections_added += 1
                    logger.info(f"Added Keywords section with {len(tag_list)} tags")
        
        # 3. Add Key People (Name, Context) - always immediately after Keywords
        if "key_people_and_tags" in results:
            key_people_tags_data = results["key_people_and_tags"]
            if isinstance(key_people_tags_data, dict):
                if "key_people" in key_people_tags_data and isinstance(key_people_tags_data["key_people"], list):
                    key_people_rows = []
                    for item in key_people_tags_data["key_people"]:
                        key_people_rows.append({
                            "Name": item.get("name", ""),
                            "Context": item.get("context", "")
                        })
                    if key_people_rows:
                        # Add 'Key People' heading
                        p = doc.add_paragraph()
                        run = p.add_run("Key People")
                        run.font.name = 'Arial'
                        run.font.size = Pt(14)
                        run.font.bold = True
                        p.space_after = Pt(6)
                        # Add table with explicit column order
                        columns = ["Name", "Context"]
                        table = doc.add_table(rows=1, cols=len(columns))
                        table.style = 'Table Grid'
                        hdr_cells = table.rows[0].cells
                        for i, col in enumerate(columns):
                            hdr_cells[i].text = col
                            for paragraph in hdr_cells[i].paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(10)
                                    run.font.bold = True
                        for row in key_people_rows:
                            row_cells = table.add_row().cells
                            for i, col in enumerate(columns):
                                row_cells[i].text = str(row.get(col, ""))
                                for paragraph in row_cells[i].paragraphs:
                                    for run in paragraph.runs:
                                        run.font.name = 'Arial'
                                        run.font.size = Pt(9)
                        doc.add_paragraph()
                        sections_added += 1
                        logger.info(f"Added Key People table with {len(key_people_rows)} entries")
        
        # 4. Add Timeline
        if "timeline_events" in results:
            timeline_data = results["timeline_events"]
            if isinstance(timeline_data, dict) and "timeline" in timeline_data:
                timeline = timeline_data["timeline"]
                if isinstance(timeline, list) and timeline:
                    add_table_section(doc, "Timeline", timeline)
                    sections_added += 1
                    logger.info(f"Added Timeline table with {len(timeline)} events")
        
        # 5. Process NER data from multiple steps
        ner_data = {}
        ner_step_names = [
            "extract_ner_people_orgs_locations",
            "extract_ner_dates_legal_rivers", 
            "extract_ner_specialized_entities"
        ]
        for step_name in ner_step_names:
            if step_name in results:
                step_data = results[step_name]
                if isinstance(step_data, dict):
                    ner_data.update(step_data)
                    logger.info(f"Added NER data from step: {step_name} (dict)")
                elif isinstance(step_data, str):
                    try:
                        parsed_step_data = json.loads(step_data)
                        if isinstance(parsed_step_data, dict):
                            ner_data.update(parsed_step_data)
                            logger.info(f"Added NER data from step: {step_name} (parsed JSON)")
                    except json.JSONDecodeError:
                        logger.warning(f"Failed to parse JSON for NER step: {step_name}")
                        continue
        
        # Add NER tables with explicit column order, no sorting
        ner_tables_added = 0
        if ner_data:
            logger.info(f"Processing NER data with {len(ner_data)} categories")
            # Persons, Organizations, Locations: Name, Alternative Spellings, Context
            for category in ["persons", "organizations", "locations"]:
                if category in ner_data and isinstance(ner_data[category], list) and ner_data[category]:
                    rows = []
                    for item in ner_data[category]:
                        rows.append({
                            "Name": item.get("name", ""),
                            "Alternative Spellings": ", ".join(item.get("alternative_spellings", [])) if isinstance(item.get("alternative_spellings", []), list) else item.get("alternative_spellings", ""),
                            "Context": item.get("context", "")
                        })
                    columns = ["Name", "Alternative Spellings", "Context"]
                    table = doc.add_table(rows=1, cols=len(columns))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for i, col in enumerate(columns):
                        hdr_cells[i].text = col
                        for paragraph in hdr_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Arial'
                                run.font.size = Pt(10)
                                run.font.bold = True
                    for row in rows:
                        row_cells = table.add_row().cells
                        for i, col in enumerate(columns):
                            row_cells[i].text = str(row.get(col, ""))
                            for paragraph in row_cells[i].paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(9)
                    doc.add_paragraph()
                    ner_tables_added += 1
                    logger.info(f"Added {category} NER table with {len(rows)} entries")
            
            # Dates, Legal Refs, Rivers: Name, Context
            for category in ["dates", "legal_refs", "rivers"]:
                if category in ner_data and isinstance(ner_data[category], list) and ner_data[category]:
                    rows = []
                    for item in ner_data[category]:
                        rows.append({
                            "Name": item.get("name", item.get("date", "")),
                            "Context": item.get("context", "")
                        })
                    columns = ["Name", "Context"]
                    table = doc.add_table(rows=1, cols=len(columns))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for i, col in enumerate(columns):
                        hdr_cells[i].text = col
                        for paragraph in hdr_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Arial'
                                run.font.size = Pt(10)
                                run.font.bold = True
                    for row in rows:
                        row_cells = table.add_row().cells
                        for i, col in enumerate(columns):
                            row_cells[i].text = str(row.get(col, ""))
                            for paragraph in row_cells[i].paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(9)
                    doc.add_paragraph()
                    ner_tables_added += 1
                    logger.info(f"Added {category} NER table with {len(rows)} entries")
            
            # All other categories: keep default behavior
            for category in ner_data:
                if category not in ["persons", "organizations", "locations", "dates", "legal_refs", "rivers"]:
                    if isinstance(ner_data[category], list) and ner_data[category]:
                        add_table_section(doc, category.replace('_', ' ').title(), ner_data[category])
                        ner_tables_added += 1
                        logger.info(f"Added {category} NER table with {len(ner_data[category])} entries")
        
        # Ensure output directory exists
        ensure_dirs(output_path)
        
        # Change extension to .docx
        output_path = output_path.with_suffix('.docx')
        
        # Save document
        doc.save(output_path)
        
        # Get relative path for output
        output_rel_path = SegmentHandler.get_relative_path(output_path)
        
        total_sections = sections_added + ner_tables_added
        logger.info(f"Successfully created Word document with {total_sections} sections")
        logger.info(f"Output saved to: {output_rel_path}")
        
        return {
            "outputs": [str(output_rel_path)],
            "source": str(rel_path),
            "success": True,
            "details": {
                "sections_created": total_sections,
                "main_sections": sections_added,
                "ner_tables": ner_tables_added,
                "output_format": "docx"
            }
        }
        
    except FileNotFoundError:
        logger.error(f"JSON file not found: {json_file_path}")
        return {
            "outputs": [],
            "source": str(rel_path),
            "success": False,
            "error": "File not found"
        }
    except json.JSONDecodeError as e:
        logger.error(f"Invalid JSON in file {json_file_path}: {e}")
        return {
            "outputs": [],
            "source": str(rel_path),
            "success": False,
            "error": f"Invalid JSON: {e}"
        }
    except Exception as e:
        logger.error(f"Error converting {json_file_path}: {e}")
        return {
            "outputs": [],
            "source": str(rel_path),
            "success": False,
            "error": str(e)
        }

def process_document(file_path: str, output_folder: Path) -> Dict:
    """Process a single JSON document file"""
    file_path = Path(file_path)
    
    def process_fn(f: Path, o: Path) -> Dict:
        return convert_json_to_word(f, o)
    
    return process_file(
        file_path=str(file_path),
        output_folder=output_folder,
        process_fn=process_fn,
        file_types={'.json': process_fn}
    )

def json_to_word(
    source_folder: Path = typer.Argument(..., help="Source folder containing JSON files"),
    source_manifest: Path = typer.Argument(..., help="Manifest file from LLM processing"),
    output_folder: Path = typer.Argument(..., help="Output folder for Word documents")
):
    """Convert JSON files from LLM processing to Word documents"""
    
    console.print(f"[blue]Converting JSON files to Word documents")
    console.print(f"[cyan]Source folder: {source_folder}")
    console.print(f"[cyan]Source manifest: {source_manifest}")
    console.print(f"[cyan]Output folder: {output_folder}")
    
    processor = BatchProcessor(
        input_manifest=source_manifest,
        output_folder=output_folder,
        process_name="json_to_word",
        base_folder=source_folder,
        processor_fn=process_document
    )
    
    result = processor.process()
    
    console.print(f"[green]Conversion completed!")
    console.print(f"[green]Processed: {result.get('processed', 0)}")
    console.print(f"[yellow]Skipped: {result.get('skipped', 0)}")
    console.print(f"[red]Failed: {result.get('failed', 0)}")

if __name__ == "__main__":
    typer.run(json_to_word) 