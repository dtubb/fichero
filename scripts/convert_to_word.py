import typer
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_BREAK
from PIL import Image
import io
from rich.console import Console
import json
import tempfile
import shutil
import subprocess

from utils.batch import BatchProcessor
from utils.processor import process_file
from utils.segment_handler import SegmentHandler
from utils.files import ensure_dirs, get_relative_path
from utils.manifest import ManifestProcessor

console = Console()
app = typer.Typer()

# Update the grey color to a lighter shade
GREY_FILL = "E8E8E8"  # Lighter grey

class SpreadManager:
    """Manages temporary storage and processing of spreads"""
    def __init__(self, temp_dir: Path):
        self.temp_dir = temp_dir
        self.temp_dir.mkdir(parents=True, exist_ok=True)
        self.spreads = {}
        
    def add_spread(self, parent_folder: str, image_path: str, text: str, filename: str) -> bool:
        """Add a spread to the temporary storage"""
        if parent_folder not in self.spreads:
            self.spreads[parent_folder] = {
                "path": None,
                "processed": False,
                "spreads": [],
                "needs_rebuild": True
            }
            
        # Check if spread exists
        for existing_spread in self.spreads[parent_folder]["spreads"]:
            if existing_spread["image_path"] == image_path:
                if existing_spread["text"] != text:
                    existing_spread["text"] = text
                    self.spreads[parent_folder]["needs_rebuild"] = True
                return True
                
        # Add new spread
        self.spreads[parent_folder]["spreads"].append({
            "image_path": image_path,
            "text": text,
            "filename": filename
        })
        self.spreads[parent_folder]["needs_rebuild"] = True
        self.spreads[parent_folder]["processed"] = True
        return True
        
    def get_spreads(self, parent_folder: str) -> list:
        """Get spreads for a parent folder"""
        return self.spreads.get(parent_folder, {}).get("spreads", [])
        
    def needs_rebuild(self, parent_folder: str) -> bool:
        """Check if a document needs rebuilding"""
        return self.spreads.get(parent_folder, {}).get("needs_rebuild", False)
        
    def is_processed(self, parent_folder: str) -> bool:
        """Check if a document has been processed"""
        return self.spreads.get(parent_folder, {}).get("processed", False)
        
    def set_path(self, parent_folder: str, path: Path):
        """Set the path for a document"""
        if parent_folder in self.spreads:
            self.spreads[parent_folder]["path"] = path
            
    def get_path(self, parent_folder: str) -> Path:
        """Get the path for a document"""
        return self.spreads.get(parent_folder, {}).get("path")
        
    def cleanup(self):
        """Clean up temporary files"""
        if self.temp_dir.exists():
            shutil.rmtree(self.temp_dir)

def set_document_properties(doc):
    """Set up initial document properties"""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.different_first_page_header_footer = True
    # No default margins - we'll set them per section

def create_cover_page(doc, folder_name):
    """Create a white cover page with centered folder name as title"""
    # Create blank first page
    blank_section = doc.sections[0]
    blank_section.left_margin = blank_section.right_margin = blank_section.top_margin = blank_section.bottom_margin = Inches(1)
    doc.add_paragraph()  # Add empty paragraph for blank page
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    
    # Create new section for cover page
    cover_section = doc.add_section(WD_SECTION.NEW_PAGE)
    cover_section.left_margin = cover_section.right_margin = cover_section.top_margin = cover_section.bottom_margin = Inches(1)
    
    # Add vertical space to center content
    for _ in range(10):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Use the full folder name as title, preserving all special characters
    title = str(folder_name)  # Ensure string conversion
    run = p.add_run(title)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(24)
    run.font.bold = True
    
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

def calculate_optimal_font_size(text_length, page_width_inches, page_height_inches):
    """Calculate optimal font size to fit text on one page"""
    # Approximate characters per line and lines per page at 12pt
    CHAR_WIDTH_12PT = 0.11
    LINE_HEIGHT_12PT = 0.18
    
    # Available space (accounting for margins)
    available_width = page_width_inches - 2
    available_height = page_height_inches - 2
    
    # Define font size ranges based on text length
    if text_length < 500:
        font_sizes = list(range(12, 10, -1))    # Short texts
    elif text_length < 1000:
        font_sizes = list(range(11, 9, -1))     # Medium texts
    else:
        font_sizes = list(range(10, 8, -1))     # Long texts
    
    for font_size in font_sizes:
        scale_factor = font_size / 12.0
        chars_per_line = int(available_width / (CHAR_WIDTH_12PT * scale_factor))
        lines_per_page = int(available_height / (LINE_HEIGHT_12PT * scale_factor))
        total_capacity = chars_per_line * lines_per_page
        
        if text_length <= total_capacity:
            return font_size
    
    return 8

def load_image_with_jxl_support(image_path: Path) -> Image.Image:
    """
    Load an image file, with support for JPEG XL format.
    Falls back to PIL for other formats.
    """
    if image_path.suffix.lower() == '.jxl':
        # Check if djxl is available
        if shutil.which('djxl'):
            try:
                # Convert JXL to temporary PNG using djxl
                temp_png = image_path.with_suffix('.temp.png')
                cmd = ['djxl', str(image_path), str(temp_png)]
                process = subprocess.run(cmd, capture_output=True, text=True)
                
                if process.returncode == 0 and temp_png.exists():
                    # Load the converted PNG
                    img = Image.open(temp_png)
                    # Clean up temp file
                    temp_png.unlink()
                    return img
                else:
                    console.print(f"[yellow]Warning: djxl conversion failed for {image_path}: {process.stderr}")
            except Exception as e:
                console.print(f"[yellow]Warning: Failed to convert JXL file {image_path}: {str(e)}")
        else:
            console.print(f"[yellow]Warning: djxl not installed, cannot read JXL file {image_path}")
            console.print("Install with: brew install libjxl (macOS) or apt-get install libjxl-tools (Ubuntu/Debian)")
        
        # If JXL conversion failed, raise an exception
        raise ValueError(f"Cannot read JXL file {image_path} - djxl not available or conversion failed")
    
    # For other formats, use PIL directly
    return Image.open(image_path)

def create_spread(doc, image_path, text, filename):
    """Create a two-page spread with image (left) and text (right)"""
    # Left page - Image
    section = doc.add_section()
    
    # Center image by setting equal margins
    margin = Inches(0.75)
    section.left_margin = section.right_margin = margin
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.75)
    
    # Add image paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    
    try:
        # Use the new function that supports JXL
        img = load_image_with_jxl_support(Path(image_path))
        
        # Convert to RGB (removing alpha channel)
        if img.mode in ('RGBA', 'LA') or (img.mode == 'P' and 'transparency' in img.info):
            bg = Image.new('RGB', img.size, 'white')
            if img.mode != 'RGBA':
                img = img.convert('RGBA')
            bg.paste(img, mask=img.split()[3])
            img = bg
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        
        # Calculate dimensions for page
        available_width = 8.5 - (2 * margin.inches)
        available_height = 9.75
        img_ratio = img.height / img.width
        
        if img_ratio > available_height/available_width:
            height = available_height
            width = height / img_ratio
        else:
            width = available_width
            height = width * img_ratio
        
        # Keep 300 DPI resolution for print quality
        target_width = int(width * 300)
        target_height = int(height * 300)
        img = img.resize((target_width, target_height), Image.Resampling.LANCZOS)
        
        # Save as JPG with compression
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format='JPEG', quality=80, optimize=True)
        img_byte_arr.seek(0)
        
        run.add_picture(img_byte_arr, width=Inches(width), height=Inches(height))
    except Exception as e:
        console.print(f"[red]Error loading image {image_path}: {str(e)}")
        run.add_text(f"[Error loading image: {str(e)}]")

    # Right page - Text
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.left_margin = section.right_margin = Inches(1)
    section.top_margin = section.bottom_margin = Inches(1)
    
    # Add filename as title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run(filename)
    title_run.font.name = 'Times New Roman'
    title_run.font.size = Pt(12)
    title_run.font.bold = True
    title_p.space_after = Pt(12)  # Add space after title
    
    # Add text content
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(calculate_optimal_font_size(len(text), 8.5, 11.0))
    
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    
    # Add page break for next spread
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    
    return True

def process_document(file_path: str, output_folder: Path, spread_manager: SpreadManager) -> dict:
    """Process a single document file using process_file utility"""
    file_path = Path(file_path)
    
    def process_fn(f: str, o: Path) -> dict:
        try:
            # Get relative path using SegmentHandler
            rel_path = SegmentHandler.get_relative_path(Path(f))
            console.print(f"\nProcessing document: {f}")
            console.print(f"Relative path: {rel_path}")
            
            # Get the parent folder name for document grouping
            parent_folder = Path(f).parent.name
            console.print(f"[blue]Parent folder name: {parent_folder}")
            
            # Find the corresponding text file using consistent path handling
            text_path = output_folder.parent / "transcriptions" / "documents" / rel_path.with_suffix('.txt')
            if not text_path.exists():
                # Try with raw path if not found
                text_path = Path(str(text_path).replace(';', '\\;'))
                if not text_path.exists():
                    console.print(f"[red]Text file not found: {text_path}")
                    return {
                        "error": f"Text file not found: {text_path}",
                        "source": str(rel_path),
                        "success": False
                    }
            
            # Read the text content
            try:
                text = text_path.read_text(encoding='utf-8')
            except UnicodeDecodeError:
                # Fallback to system default encoding if UTF-8 fails
                text = text_path.read_text()
            
            # Add spread to manager
            spread_manager.add_spread(parent_folder, str(f), text, Path(f).stem)
            
            # Get the final document path using consistent path handling
            doc_output_path = output_folder / "documents" / rel_path.parent / f"{parent_folder}.docx"
            spread_manager.set_path(parent_folder, doc_output_path)
            
            return {
                "outputs": [str(SegmentHandler.get_relative_path(doc_output_path))],
                "source": str(rel_path),
                "details": {
                    "text_length": len(text),
                    "needs_rebuild": spread_manager.needs_rebuild(parent_folder),
                    "spread_count": len(spread_manager.get_spreads(parent_folder))
                }
            }
            
        except Exception as e:
            console.print(f"[red]Error processing {f}: {str(e)}")
            return {
                "error": str(e),
                "source": str(rel_path)
            }
    
    return process_file(
        file_path=str(file_path),
        output_folder=output_folder,
        process_fn=lambda f, o: process_fn(f, o),
        file_types={
            '.png': process_fn,
            '.jpg': process_fn,
            '.jpeg': process_fn,
            '.jxl': process_fn  # Add JXL support
        }
    )

@app.command()
def convert_to_word(
    background_removed_folder: Path = typer.Argument(..., help="Input background removed images folder"),
    transcription_manifest: Path = typer.Argument(..., help="Input transcription manifest"),
    word_folder: Path = typer.Argument(..., help="Output folder for Word documents")
):
    """Convert background-removed images and transcriptions to Word documents with side-by-side layout"""
    console.print(f"[green]Converting images in {background_removed_folder} to Word documents")
    
    # Create temporary directory for spreads
    with tempfile.TemporaryDirectory() as temp_dir:
        spread_manager = SpreadManager(Path(temp_dir))
        
        # Create processor for PNG files
        processor = BatchProcessor(
            input_manifest=transcription_manifest,
            output_folder=word_folder,
            process_name="convert_to_word",
            base_folder=background_removed_folder / "documents",  # Ensure documents/ is included
            processor_fn=lambda f, o: process_document(f, o, spread_manager),
            use_source=True  # Use source paths from manifest
        )
        
        results = processor.process()
        
        # Process accumulated spreads and save complete documents
        for parent_folder in spread_manager.spreads:
            if not spread_manager.is_processed(parent_folder):
                console.print(f"[yellow]Warning: Document for folder {parent_folder} was not processed")
                continue
                
            try:
                # Only rebuild if needed
                if spread_manager.needs_rebuild(parent_folder):
                    # Create new document
                    doc = Document()
                    set_document_properties(doc)
                    create_cover_page(doc, parent_folder)
                    
                    # Add all spreads to the document
                    for spread in spread_manager.get_spreads(parent_folder):
                        create_spread(doc, spread["image_path"], spread["text"], spread["filename"])
                    
                    # Save complete document using consistent path handling
                    out_path = spread_manager.get_path(parent_folder)
                    ensure_dirs(out_path)
                    doc.save(str(out_path))
                    console.print(f"[green]Saved complete document: {out_path}")
                else:
                    console.print(f"[yellow]Skipping unchanged document: {parent_folder}")
            except Exception as e:
                console.print(f"[red]Error saving document for {parent_folder}: {str(e)}")
        
        return results

if __name__ == "__main__":
    app()